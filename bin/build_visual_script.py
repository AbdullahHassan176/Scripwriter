#!/usr/bin/env python3
"""
Visual Script Builder

Parses a script markdown, fetches matching images from Wikimedia Commons
for archival/footage markers, and generates two output files:
  1. A self-contained HTML — rich visual read-through in any browser
  2. A DOCX — editable Word document with inline visuals

Usage:
    python bin/build_visual_script.py "generatedScripts/SA Captured Rainbow/dark_side_of_south_africa_script.md"
    python bin/build_visual_script.py SCRIPT.md --package PRODUCTION_PACKAGE.md
    python bin/build_visual_script.py SCRIPT.md --no-fetch   # skip Wikimedia (faster)
    python bin/build_visual_script.py SCRIPT.md --html-only  # skip DOCX generation

Outputs (next to the script file):
    visual_script.html
    visual_script.docx
"""

import argparse
import base64
import html as html_lib
import io
import json
import re
import sys
import time
from dataclasses import dataclass, field
from pathlib import Path
from typing import Optional
from urllib.parse import quote

import math

import urllib3
import requests
import requests.packages.urllib3
urllib3.disable_warnings(urllib3.exceptions.InsecureRequestWarning)
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

# ─── Configurable constants ───────────────────────────────────────────────────

COMMONS_SEARCH = "https://commons.wikimedia.org/w/api.php"
WIKIPEDIA_API  = "https://en.wikipedia.org/w/api.php"
WIKIMEDIA_THUMB_WIDTH = 500
REQUEST_DELAY = 1.5   # seconds between API calls
REQUEST_TIMEOUT = 12
REQUEST_RETRY_DELAY = 15  # seconds to wait after a 429

# Keywords in SHOW descriptions that imply CUSTOM (no image to find)
CUSTOM_KEYWORDS = [
    "data card", "title card", "quote card", "section title", "talking head",
    "graphic", "org chart", "bar chart", "pie chart", "dial", "icon grid",
    "split screen", "animation", "animated", "motion graphic", "ae comp",
    "two-column", "lower third", "end card", "subscribe", "score",
    "comparative table", "timeline graphic", "icon", "icons",
    "revolving door", "line chart", "trend", "budget vs actual",
]

SELF_KEYWORDS = ["talking head", "on camera", "on-camera", "self", "hold on"]

ARCHIVE_KEYWORDS = ["archival", "footage", "b-roll", "b roll", "clip", "video"]

GE_KEYWORDS = ["aerial", "google earth", "pullback", "pull back", "satellite"]

# Noise words to strip when building a Wikimedia search query
QUERY_NOISE = re.compile(
    r"\b(archival|footage|b-roll|b roll|aerial|pullback|pull back|slow|wide|"
    r"close-up|closeup|widens|fade|dissolve|hold|cut to|quick|shot|appears|"
    r"starts|begins|dissolves|reveal|zoom|tight|label|animate|animated|"
    r"show|left|right|center|above|below|white|black|dark|background|"
    r"begins|same|final|image|seconds|card|graphic|dissolve|score|drops|"
    r"the|and|or|to|in|on|at|of|a|an|is|with|from|for|its|it|this|that|"
    r"every|all|its|as|by|but|into|over|up|no|not)\b",
    re.IGNORECASE,
)

# File types Wikimedia may return that aren't useful images
SKIP_EXTENSIONS = (".pdf", ".svg", ".ogg", ".ogv", ".webm", ".mp3", ".flac")


# ─── Data classes ─────────────────────────────────────────────────────────────

@dataclass
class ImageResult:
    url: str = ""
    thumb_url: str = ""
    title: str = ""
    author: str = ""
    license: str = ""
    commons_url: str = ""
    description: str = ""
    error: str = ""

    @property
    def found(self) -> bool:
        return bool(self.thumb_url)


@dataclass
class Segment:
    kind: str  # section_header | visual_cue | text | pause | divider | quote | note
    content: str
    # only for visual_cue:
    source_type: str = ""   # ARCHIVE | CUSTOM | SELF | GE | LICENSED | GENERAL
    image: ImageResult = field(default_factory=ImageResult)
    prod_source: str = ""   # from production package (AP, W, CUSTOM, etc.)
    prod_notes: str = ""    # notes from production package
    shot_num: str = ""
    # only for text:
    delivery: str = "VO"    # VO | TH | SPONSOR | BUMP


# ─── Script parsing ───────────────────────────────────────────────────────────

SHOW_RE = re.compile(r"\[SHOW:\s*(.*?)\]", re.IGNORECASE | re.DOTALL)
SHOW_LINE_RE = re.compile(r"^\[SHOW:\s*(.+?)\]\s*$", re.IGNORECASE)
DELIVERY_PREFIX_RE = re.compile(r"^\[(VO|TH|SP|SUB)\]\s*", re.IGNORECASE)
PAUSE_RE = re.compile(r"\[PAUSE\]", re.IGNORECASE)
SECTION_RE = re.compile(r"^#{1,4}\s*(?:\[([^\]]+)\]|(.+))")
DOCX_SECTION_FILL = "1a1a1a"
DIRECTOR_NOTE_RE = re.compile(r"^#{1,4}\s*(director|style|visual convention|blueprint)", re.IGNORECASE)
HEADER_BLOCK_RE = re.compile(r"^[━=]{5,}")


def classify_visual(desc: str) -> str:
    d = desc.lower()
    if any(k in d for k in SELF_KEYWORDS):
        return "SELF"
    if any(k in d for k in CUSTOM_KEYWORDS):
        return "CUSTOM"
    if any(k in d for k in GE_KEYWORDS):
        return "GE"
    if any(k in d for k in ARCHIVE_KEYWORDS):
        return "ARCHIVE"
    return "GENERAL"


def extract_search_query(desc: str) -> str:
    """
    Turn a SHOW description into a focused Wikimedia search string.
    Prioritises the content after a dash (the actual subject),
    keeps capitalised/named terms, and drops technical stage-direction words.
    Returns empty string if the result is too vague to be useful.
    """
    # If there's a dash separator, take the part after it (the subject)
    if "—" in desc:
        candidate = desc.split("—", 1)[1]
    elif " — " in desc:
        candidate = desc.split(" — ", 1)[1]
    else:
        candidate = desc

    # Strip parenthetical asides like "(left)" or "(Zuma years highlighted)"
    candidate = re.sub(r"\([^)]*\)", "", candidate)
    # Strip things after a period (often camera notes)
    candidate = candidate.split(".")[0]
    # Take content before the first slash (split-screen usually describes first half well)
    candidate = candidate.split("/")[0].strip()

    # Remove noise words
    candidate = QUERY_NOISE.sub(" ", candidate)
    # Strip leftover punctuation
    candidate = re.sub(r"[^\w\s]", " ", candidate)
    candidate = re.sub(r"\s+", " ", candidate).strip()

    words = candidate.split()
    # Drop pure-numeric tokens that don't give context (e.g. "73", "1996" alone)
    meaningful = [w for w in words if not re.match(r"^\d+$", w) or len(w) == 4]
    # Cap at 5 words
    useful = meaningful[:5]

    # Must have at least 2 meaningful words to be worth searching
    if len(useful) < 2:
        return ""
    return " ".join(useful)


def _find_script_body_range(lines: list[str]) -> tuple[int, int]:
    """
    Return (start, end) indices of the actual script body.
    - Skips the opening ━━━ … ━━━ metadata header AND any preamble text
      (e.g. VISUAL CONVENTION section) by starting at the first ## section.
    - Ends before the first ━━━ that comes after the last ## section
      (director notes / data sources appendix).
    """
    # Find the second ━━━ (end of opening header block)
    after_header = 0
    count = 0
    for i, line in enumerate(lines):
        if HEADER_BLOCK_RE.match(line):
            count += 1
            if count == 2:
                after_header = i + 1
                break

    # Start = first ## section header (skips VISUAL CONVENTION preamble)
    start = after_header
    for i in range(after_header, len(lines)):
        if SECTION_RE.match(lines[i]):
            start = i
            break

    # End = first ━━━ after the last ## section header
    last_section_idx = start
    for i in range(start, len(lines)):
        if SECTION_RE.match(lines[i]):
            last_section_idx = i
    end = len(lines)
    for i in range(last_section_idx, len(lines)):
        if HEADER_BLOCK_RE.match(lines[i]):
            end = i
            break

    return start, end


def parse_script(path: Path) -> list[Segment]:
    text = path.read_text(encoding="utf-8")
    lines = text.splitlines()
    start, end = _find_script_body_range(lines)
    body_lines = lines[start:end]

    segments: list[Segment] = []
    buffer: list[str] = []

    def flush_buffer():
        nonlocal buffer
        chunk = "\n".join(buffer).strip()
        if chunk:
            stripped_inner = chunk.strip("*").strip()
            if chunk.startswith("*") and chunk.endswith("*") and "\n" not in chunk:
                segments.append(Segment(kind="quote", content=stripped_inner))
            else:
                segments.append(Segment(kind="text", content=chunk))
        buffer = []

    i = 0
    while i < len(body_lines):
        line = body_lines[i]

        # Dividers
        if re.match(r"^-{3,}$", line.strip()):
            flush_buffer()
            segments.append(Segment(kind="divider", content=""))
            i += 1
            continue

        # Section headers (## [HOOK], ## [SECTION 1 — ...])
        m = SECTION_RE.match(line)
        if m:
            flush_buffer()
            title = (m.group(1) or m.group(2) or "").strip()
            if title:
                segments.append(Segment(kind="section_header", content=title))
            i += 1
            continue

        # Pause
        if PAUSE_RE.search(line):
            flush_buffer()
            segments.append(Segment(kind="pause", content=""))
            i += 1
            continue

        # SHOW marker
        show_m = SHOW_RE.search(line)
        if show_m:
            before = line[: show_m.start()].strip()
            if before:
                buffer.append(before)
            flush_buffer()
            desc = show_m.group(1).strip()
            src_type = classify_visual(desc)
            segments.append(Segment(kind="visual_cue", content=desc, source_type=src_type))
            after = line[show_m.end():].strip()
            if after:
                buffer.append(after)
            i += 1
            continue

        # Regular line
        buffer.append(line)
        i += 1

    flush_buffer()
    return [s for s in segments if s.content or s.kind in ("divider", "pause")]


# ─── Production package parsing ───────────────────────────────────────────────

TABLE_ROW_RE = re.compile(r"^\|\s*(.*?)\s*\|")


def parse_production_package(path: Path) -> list[dict]:
    """Returns list of {shot_num, script_line, visual, source, notes}."""
    text = path.read_text(encoding="utf-8")
    rows = []
    header_seen = False
    for line in text.splitlines():
        if not line.startswith("|"):
            header_seen = False
            continue
        cells = [c.strip() for c in line.strip("|").split("|")]
        if len(cells) < 3:
            continue
        # Skip separator rows
        if all(re.match(r"^-+$", c) for c in cells if c):
            header_seen = True
            continue
        # Skip header rows
        if not header_seen:
            continue
        if len(cells) >= 5 and cells[0].isdigit():
            rows.append({
                "shot_num": cells[0],
                "script_line": cells[1].strip('"'),
                "visual": cells[2],
                "source": cells[3],
                "notes": cells[4] if len(cells) > 4 else "",
            })
    return rows


_DELIVERY_FROM_PREFIX = {"VO": "VO", "TH": "TH", "SP": "SPONSOR", "SUB": "BUMP"}


def assign_delivery_types(segments: list[Segment], *, only_unset: bool = False):
    """
    Post-parse pass: set delivery type on every text segment.
    - Default: VO (voiceover over B-roll / graphics)
    - After a SELF visual cue: next text block is TH (talking head, on camera)
    - Sponsor/subscribe text detected by keyword

    If only_unset=True, leaves segments that already have a non-default delivery
    from an explicit [VO]/[TH]/[SP]/[SUB] prefix in the DOCX.
    """
    pending = "VO"
    for seg in segments:
        if seg.kind == "visual_cue":
            pending = "TH" if seg.source_type == "SELF" else "VO"
        elif seg.kind == "text":
            if only_unset and seg.delivery not in ("", "VO"):
                pending = "VO"
                continue
            txt = seg.content.lower()
            if any(k in txt for k in ["[sponsor", "quick word from", "code 2and20",
                                       "[offer]", "sponsor read", "this week's sponsor"]):
                seg.delivery = "SPONSOR"
            elif any(k in txt for k in ["96% of you", "hit subscribe", "aren't subscribed",
                                         "back to the video"]):
                seg.delivery = "BUMP"
            else:
                seg.delivery = pending
                if pending == "TH":   # TH is sticky for one block then resets
                    pending = "VO"


# ─── DOCX round-trip (edit Word → sync script + HTML) ─────────────────────────

def _paragraph_shading_fill(paragraph) -> str:
    """Return hex fill colour of paragraph shading, or ''."""
    p_pr = paragraph._p.pPr
    if p_pr is None:
        return ""
    shd = p_pr.find(qn("w:shd"))
    if shd is None:
        return ""
    return (shd.get(qn("w:fill")) or "").lower()


def _iter_docx_blocks(document):
    """Yield paragraphs and tables in document order."""
    from docx.oxml.table import CT_Tbl
    from docx.oxml.text.paragraph import CT_P
    from docx.table import Table
    from docx.text.paragraph import Paragraph

    parent = document.element.body
    for child in parent.iterchildren():
        if isinstance(child, CT_P):
            yield Paragraph(child, document)
        elif isinstance(child, CT_Tbl):
            yield Table(child, document)


def _is_docx_section_header(paragraph, text: str) -> bool:
    if _paragraph_shading_fill(paragraph) == DOCX_SECTION_FILL:
        return True
    if text.startswith("## "):
        return True
    return False


def parse_visual_docx(path: Path) -> list[Segment]:
    """
    Parse visual_script.docx back into segments.
    Expects the round-trip format from build_docx:
      - Section headers (dark shaded paragraph)
      - [VO] / [TH] / [SP] / [SUB] prefixed narration
      - [SHOW: description] on its own line
      - [PAUSE], ---, quoted blocks
    """
    doc = Document(str(path))
    segments: list[Segment] = []
    text_parts: list[str] = []
    current_delivery = "VO"
    delivery_from_prefix = False

    def flush_text():
        nonlocal text_parts, current_delivery, delivery_from_prefix
        if not text_parts:
            return
        content = "\n\n".join(text_parts).strip()
        if content:
            seg = Segment(kind="text", content=content, delivery=current_delivery)
            segments.append(seg)
        text_parts = []
        delivery_from_prefix = False

    skip_next_table = False

    for block in _iter_docx_blocks(doc):
        if hasattr(block, "rows"):  # Table — image attachment for previous SHOW
            if segments and segments[-1].kind == "visual_cue":
                continue
            continue

        text = block.text.strip()
        if not text:
            continue

        # Skip title / meta (first lines before first section)
        if not segments and not _is_docx_section_header(block, text):
            if text.startswith("~") and "words" in text:
                continue
            if len(text) < 80 and block.runs and block.runs[0].bold and block.runs[0].font.size and block.runs[0].font.size.pt >= 16:
                continue

        if _is_docx_section_header(block, text):
            flush_text()
            title = text.lstrip("#").strip().strip("[]")
            segments.append(Segment(kind="section_header", content=title))
            current_delivery = "VO"
            continue

        if text in ("[ PAUSE ]", "[PAUSE]") or PAUSE_RE.fullmatch(text):
            flush_text()
            segments.append(Segment(kind="pause", content=""))
            continue

        if re.match(r"^-{3,}$", text) or text == "─" * 40:
            flush_text()
            segments.append(Segment(kind="divider", content=""))
            continue

        show_m = SHOW_LINE_RE.match(text) or SHOW_RE.fullmatch(text)
        if show_m:
            flush_text()
            desc = show_m.group(1).strip()
            segments.append(Segment(
                kind="visual_cue",
                content=desc,
                source_type=classify_visual(desc),
            ))
            skip_next_table = True
            continue

        dm = DELIVERY_PREFIX_RE.match(text)
        if dm:
            flush_text()
            current_delivery = _DELIVERY_FROM_PREFIX.get(dm.group(1).upper(), "VO")
            delivery_from_prefix = True
            rest = text[dm.end():].strip()
            if rest:
                text_parts.append(rest)
            continue

        if text.startswith('"') and text.endswith('"') and len(text) > 2:
            flush_text()
            segments.append(Segment(kind="quote", content=text[1:-1].strip()))
            continue

        # Continuation paragraph (no prefix) — same spoken block
        text_parts.append(text)

    flush_text()
    return [s for s in segments if s.content or s.kind in ("divider", "pause")]


def segments_to_script_body(segments: list[Segment]) -> str:
    """Serialize segments to markdown script body (between header/footer blocks)."""
    lines: list[str] = []
    for seg in segments:
        if seg.kind == "section_header":
            if lines:
                lines.append("")
            lines.append(f"## [{seg.content}]")
            lines.append("")
        elif seg.kind == "text":
            lines.append(seg.content)
            lines.append("")
        elif seg.kind == "visual_cue":
            lines.append(f"[SHOW: {seg.content}]")
            lines.append("")
        elif seg.kind == "pause":
            lines.append("[PAUSE]")
            lines.append("")
        elif seg.kind == "quote":
            lines.append(f"*{seg.content}*")
            lines.append("")
        elif seg.kind == "divider":
            lines.append("---")
            lines.append("")
    while lines and lines[-1] == "":
        lines.pop()
    return "\n".join(lines)


def update_script_from_segments(script_path: Path, segments: list[Segment]) -> None:
    """Replace the script body in the .md file, keeping header/footer blocks."""
    text = script_path.read_text(encoding="utf-8")
    file_lines = text.splitlines()
    start, end = _find_script_body_range(file_lines)
    new_body = segments_to_script_body(segments).splitlines()
    updated = file_lines[:start] + new_body + file_lines[end:]
    out = "\n".join(updated)
    if not out.endswith("\n"):
        out += "\n"
    script_path.write_text(out, encoding="utf-8")


def sync_from_docx(
    docx_path: Path,
    script_path: Path,
    package_path: Optional[Path] = None,
    *,
    no_fetch: bool = False,
    html_only: bool = False,
    docx_only: bool = False,
    context: str = "",
) -> None:
    """Read edited visual_script.docx → update .md → regenerate HTML (+ DOCX)."""
    print(f"\n=== Sync from DOCX ===")
    print(f"DOCX   : {docx_path.name}")
    print(f"Script : {script_path.name}")

    print("Parsing DOCX...", end=" ")
    segments = parse_visual_docx(docx_path)
    visual_count = sum(1 for s in segments if s.kind == "visual_cue")
    print(f"OK  {len(segments)} segments, {visual_count} visual cues")

    assign_delivery_types(segments, only_unset=True)

    if package_path and package_path.exists():
        print("Matching production package...", end=" ")
        rows = parse_production_package(package_path)
        match_visuals_to_package(segments, rows)
        matched = sum(1 for s in segments if s.kind == "visual_cue" and s.prod_source)
        print(f"OK  {matched}/{visual_count} matched")

    if not no_fetch:
        print("\nFetching images...")
        stem = script_path.stem.replace("_", " ").replace("-", " ")
        context_hint = context or stem[:30]
        fetch_images_for_segments(segments, verbose=True, context_hint=context_hint)

    print("Updating script markdown...", end=" ")
    update_script_from_segments(script_path, segments)
    print("OK")

    out_dir = script_path.parent
    html_out = out_dir / "visual_script.html"
    docx_out = out_dir / "visual_script.docx"

    if not docx_only:
        print("Building HTML...", end=" ")
        html_content = build_html(segments, script_path, package_path)
        html_out.write_text(html_content, encoding="utf-8")
        print(f"OK  {html_out.name}")

    if not html_only:
        print("Building DOCX...", end=" ")
        doc = build_docx(segments, script_path)
        doc.save(str(docx_out))
        print(f"OK  {docx_out.name}")

    print(f"\nDone! Edit {docx_out.name} in Word, then run again with --sync-docx to push changes.")


def match_visuals_to_package(segments: list[Segment], package_rows: list[dict]):
    """Best-effort: assign prod_source and prod_notes to visual_cue segments."""
    visual_segs = [s for s in segments if s.kind == "visual_cue"]
    # Filter out SELF rows from the package for matching purposes
    matchable_rows = [r for r in package_rows if "SELF" not in r["source"].upper()]

    def word_overlap(a: str, b: str) -> int:
        wa = set(re.findall(r"\w+", a.lower()))
        wb = set(re.findall(r"\w+", b.lower()))
        return len(wa & wb)

    used: set[int] = set()
    for seg in visual_segs:
        best_score = -1
        best_idx = -1
        for idx, row in enumerate(matchable_rows):
            if idx in used:
                continue
            score = word_overlap(seg.content, row["visual"]) + word_overlap(seg.content, row["script_line"])
            if score > best_score:
                best_score = score
                best_idx = idx
        if best_idx >= 0 and best_score > 1:
            row = matchable_rows[best_idx]
            seg.prod_source = row["source"]
            seg.prod_notes = row["notes"]
            seg.shot_num = row["shot_num"]
            used.add(best_idx)


# ─── Wikimedia Commons image fetching ─────────────────────────────────────────

_session = requests.Session()
_session.headers["User-Agent"] = "ScriptVisualBuilder/1.0 (youtube-scriptwriter-tool)"
_session.verify = False  # Windows SSL chain quirk workaround


def _wiki_get(base_url: str, **params) -> dict:
    params["format"] = "json"
    try:
        r = _session.get(base_url, params=params, timeout=REQUEST_TIMEOUT)
        r.raise_for_status()
        return r.json()
    except Exception as e:
        return {"error": str(e)}


def _commons_api(**params) -> dict:
    return _wiki_get(COMMONS_SEARCH, **params)


def _wikipedia_api(**params) -> dict:
    return _wiki_get(WIKIPEDIA_API, **params)


def _with_retry(fn, *args, **kwargs) -> dict:
    """Call fn(*args, **kwargs), retry once after delay on 429."""
    result = fn(*args, **kwargs)
    if "error" in result and "429" in str(result.get("error", "")):
        time.sleep(REQUEST_RETRY_DELAY)
        result = fn(*args, **kwargs)
    return result


def _search_wikipedia_article_image(query: str) -> ImageResult:
    """
    Search for a Wikipedia article about the query and return its main image.
    This gives highly relevant, curated images for specific topics.
    """
    # Search for the best-matching article
    data = _with_retry(_wikipedia_api, action="query", list="search",
                       srsearch=query, srlimit=3)
    results = data.get("query", {}).get("search", [])
    if not results:
        return ImageResult(error=f"no article: '{query}'")

    page_title = results[0]["title"]

    # Fetch the article thumbnail
    time.sleep(REQUEST_DELAY)
    info = _with_retry(_wikipedia_api, action="query", titles=page_title,
                       prop="pageimages|pageprops", pithumbsize=WIKIMEDIA_THUMB_WIDTH,
                       pilicense="any")
    pages = info.get("query", {}).get("pages", {})
    page = next(iter(pages.values()), {})
    thumbnail = page.get("thumbnail", {})

    if thumbnail and thumbnail.get("source"):
        wp_url = f"https://en.wikipedia.org/wiki/{quote(page_title.replace(' ', '_'))}"
        return ImageResult(
            url=thumbnail["source"],
            thumb_url=thumbnail["source"],
            title=page_title,
            author="Wikipedia",
            license="Wikipedia (see article)",
            commons_url=wp_url,
            description=f"Main image from Wikipedia: {page_title}",
        )
    return ImageResult(error=f"no thumbnail for article '{page_title}'")


def _search_commons_file(query: str) -> ImageResult:
    """
    Search Wikimedia Commons for an image file matching the query.
    Skips non-image file types.
    """
    data = _with_retry(_commons_api, action="query", list="search", srsearch=query,
                       srnamespace=6, srlimit=8)
    if "error" in data and "query" not in data:
        return ImageResult(error=data["error"])

    results = data.get("query", {}).get("search", [])
    if not results:
        return ImageResult(error=f"no results for '{query}'")

    # Prefer photo files over PDFs/SVGs
    file_title = None
    for r in results:
        t = r["title"].lower()
        if not any(t.endswith(ext) for ext in SKIP_EXTENSIONS):
            file_title = r["title"]
            break
    if file_title is None:
        return ImageResult(error=f"only non-image results for '{query}'")

    time.sleep(REQUEST_DELAY)
    info_data = _with_retry(_commons_api, action="query", titles=file_title,
                            prop="imageinfo", iiprop="url|thumburl|extmetadata",
                            iiurlwidth=WIKIMEDIA_THUMB_WIDTH)
    pages = info_data.get("query", {}).get("pages", {})
    page = next(iter(pages.values()), {})
    ii = (page.get("imageinfo") or [{}])[0]

    thumb = ii.get("thumburl") or ii.get("url", "")
    full_url = ii.get("url", "")
    meta = ii.get("extmetadata", {})

    def _m(key: str) -> str:
        return re.sub(r"<[^>]+>", "", meta.get(key, {}).get("value", "")).strip()

    if not thumb:
        return ImageResult(error=f"no thumbnail for '{file_title}'")

    return ImageResult(
        url=full_url, thumb_url=thumb,
        title=file_title.replace("File:", ""),
        author=(_m("Artist") or _m("Credit") or "Wikimedia Commons")[:80],
        license=_m("LicenseShortName") or _m("License") or "See Commons",
        commons_url=f"https://commons.wikimedia.org/wiki/{quote(file_title)}",
        description=_m("ImageDescription")[:120],
    )


def _result_is_relevant(query: str, result_title: str) -> bool:
    """
    Basic relevance check: at least one meaningful word (4+ chars) from the
    query must appear in the result title. Prevents wildly off-topic matches.
    """
    q_words = set(re.findall(r"\b[a-zA-Z]{4,}\b", query.lower()))
    r_words = set(re.findall(r"\b[a-zA-Z]{4,}\b", result_title.lower()))
    return bool(q_words & r_words)


def search_wikimedia(query: str) -> ImageResult:
    """
    Try Wikipedia article image first (most relevant), then fall back to
    Commons file search. Filters out off-topic results.
    """
    if not query.strip() or len(query.split()) < 2:
        return ImageResult(error="query too short")

    # 1. Try Wikipedia article image
    result = _search_wikipedia_article_image(query)
    if result.found:
        if _result_is_relevant(query, result.title):
            return result
        # Wikipedia gave an off-topic article — try Commons instead
        result = ImageResult(error=f"Wikipedia result not relevant: '{result.title}'")

    # 2. Fall back to Commons file search
    time.sleep(REQUEST_DELAY)
    commons = _search_commons_file(query)
    if commons.found and _result_is_relevant(query, commons.title):
        return commons
    if commons.found:
        return ImageResult(error=f"Commons result not relevant: '{commons.title}'")
    return commons


def fetch_image_bytes(url: str) -> bytes:
    try:
        r = _session.get(url, timeout=REQUEST_TIMEOUT)
        r.raise_for_status()
        return r.content
    except Exception:
        return b""


def fetch_images_for_segments(segments: list[Segment], verbose: bool = True,
                               context_hint: str = ""):
    """
    Populate segment.image for visual_cues where appropriate.
    context_hint: optional topic context to append to queries (e.g. 'South Africa').
    """
    fetchable = [
        s for s in segments
        if s.kind == "visual_cue" and s.source_type not in ("CUSTOM", "SELF")
    ]
    total = len(fetchable)
    done = 0
    for seg in fetchable:
        query = extract_search_query(seg.content)
        if not query:
            continue

        # Append context hint if the query doesn't already contain it
        if context_hint and context_hint.lower() not in query.lower():
            # Only add hint if query has some meaningful content
            query = f"{query} {context_hint}"

        done += 1
        if verbose:
            print(f"  [{done}/{total}] Searching: '{query}'...", end=" ", flush=True)

        result = search_wikimedia(query)
        if result.found:
            if verbose:
                print(f"OK  {result.title[:55]}")
        else:
            if verbose:
                print(f"NO  {result.error[:60]}")

        seg.image = result
        time.sleep(REQUEST_DELAY)


# ─── Data chart parsing & SVG rendering ──────────────────────────────────────

def _parse_chart_data(desc: str) -> Optional[dict]:
    """Extract structured chart data from a SHOW description. Returns None if unrecognised."""
    d = desc

    # Triple gauge / dials: 'Official: 33%" / "Expanded: 43%" / "Ages 15–24: 61%"'
    triples = re.findall(r'"?([A-Za-z][A-Za-z0-9\s\u2013\-]+?)"?\s*:\s*(\d+)%', d)
    vals = [(lbl.strip(), int(v)) for lbl, v in triples if 0 < int(v) <= 100]
    if len(vals) >= 3:
        return {"type": "gauges", "values": vals[:3]}

    # Trend line: "1994: 43% → 2010: 50% → 2024: 61%"
    trend = re.findall(r"(\d{4})\s*:\s*(\d+)%", d)
    if len(trend) >= 2:
        return {"type": "line", "points": [(int(yr), int(v)) for yr, v in trend]}

    # Flat trend (Gini): "nearly flat line" + year range
    if "flat" in d.lower():
        years = [int(y) for y in re.findall(r"\b(19\d{2}|20\d{2})\b", d)]
        if len(years) >= 2:
            y1, y2 = min(years), max(years)
            mid = (y1 + y2) // 2
            return {"type": "line", "flat": True,
                    "points": [(y1, 63), (y1+6, 62), (mid, 63), (y2-6, 64), (y2, 63)]}

    # Comparison bar (target vs actual, budget vs actual)
    lm = re.search(r'(?:target|budget)\D{0,10}?(?:R)?(\d+)([%bn]*)', d, re.I)
    rm = re.search(r'(?:actual|cost|over)\D{0,10}?(?:Over\s+)?(?:R)?(\d+)([%bn]*)', d, re.I)
    if lm and rm:
        lv, rv = int(lm.group(1)), int(rm.group(1))
        lu, ru = lm.group(2) or "", rm.group(2) or ""
        return {"type": "comparison_bar",
                "left": lv, "left_label": f"{lv}{lu}",
                "right": rv, "right_label": f"{rv}{ru}+"}

    # Person / icon grid: "100 people, 61 greyed"
    gm = re.search(r"(\d+)\s+(?:people|icons?).*?(\d+)\s+grey", d, re.I)
    if gm:
        return {"type": "person_grid", "total": int(gm.group(1)), "hi": int(gm.group(2))}

    # Wealth ratio: "8x the wealth"
    wm = re.search(r"(\d+)x\s+the\s+(?:wealth|income)", d, re.I)
    if wm:
        return {"type": "ratio_bar", "ratio": int(wm.group(1))}

    # Checklist / split card
    cks = re.search(r"checkmarks?\s*:\s*(.+?)(?:question marks?|\Z)", d, re.I | re.DOTALL)
    xks = re.search(r"question marks?\s*:\s*(.+?)(?:\Z)", d, re.I | re.DOTALL)
    if cks or xks:
        def _items(text: str) -> list[str]:
            # extract quoted items or colon-separated labels
            found = re.findall(r'"([^"]+)"', text)
            if not found:
                found = re.findall(r"([A-Z][a-z][^.,;\"]+?)(?:[.,;]|$)", text)
            return [f.strip() for f in found if f.strip()][:4]
        good = _items(cks.group(1) if cks else "")
        bad  = _items(xks.group(1) if xks else "")
        if good or bad:
            return {"type": "checklist", "good": good, "bad": bad}

    return None


def _svg_wrap(w: int, h: int, content: str) -> str:
    return (f'<svg viewBox="0 0 {w} {h}" width="{w}" height="{h}" '
            f'xmlns="http://www.w3.org/2000/svg" style="display:block;border-radius:4px">'
            f'{content}</svg>')


def _render_gauges(values: list[tuple[str, int]]) -> str:
    n = min(len(values), 3)
    W, H = 240, 110
    gw = W // n
    COLORS = ["#1a6ea8", "#e67e22", "#c0392b"]
    parts = [f'<rect width="{W}" height="{H}" fill="#fafafa" rx="4"/>']
    for i, (label, value) in enumerate(values[:n]):
        cx = gw * i + gw // 2
        cy, r = 72, 30
        # Background semicircle (sweep counterclockwise through top)
        parts.append(
            f'<path d="M {cx-r} {cy} A {r} {r} 0 0 0 {cx+r} {cy}" '
            f'fill="none" stroke="#e0e0e0" stroke-width="7" stroke-linecap="round"/>')
        pct = min(value / 100, 0.999)
        angle = math.pi * (1 - pct)
        ex, ey = cx + r * math.cos(angle), cy - r * math.sin(angle)
        large = 1 if pct > 0.5 else 0
        c = COLORS[i % len(COLORS)]
        if value > 0:
            parts.append(
                f'<path d="M {cx-r} {cy} A {r} {r} 0 {large} 0 {ex:.1f} {ey:.1f}" '
                f'fill="none" stroke="{c}" stroke-width="7" stroke-linecap="round"/>')
        parts.append(f'<text x="{cx}" y="{cy+3}" text-anchor="middle" '
                     f'font-size="13" font-weight="bold" fill="{c}" font-family="sans-serif">{value}%</text>')
        parts.append(f'<text x="{cx}" y="{cy+17}" text-anchor="middle" '
                     f'font-size="8" fill="#888" font-family="sans-serif">{label[:14]}</text>')
    return _svg_wrap(W, H, "".join(parts))


def _render_line_chart(points: list[tuple[int, int]], flat: bool = False) -> str:
    W, H = 240, 115
    PL, PR, PT, PB = 28, 8, 12, 22
    xs = [p[0] for p in points]
    ys = [p[1] for p in points]
    xmin, xmax = min(xs), max(xs)
    ymin, ymax = max(0, min(ys) - 8), min(100, max(ys) + 8)
    if ymax == ymin:
        ymax = ymin + 10

    def px(x): return PL + (x - xmin) / (xmax - xmin) * (W - PL - PR)
    def py(y): return H - PB - (y - ymin) / (ymax - ymin) * (H - PT - PB)

    parts = [f'<rect width="{W}" height="{H}" fill="#fafafa" rx="4"/>']
    # Horizontal grid lines
    for yt in range(0, 101, 20):
        if ymin - 2 <= yt <= ymax + 2:
            yp = py(yt)
            parts.append(f'<line x1="{PL}" y1="{yp:.1f}" x2="{W-PR}" y2="{yp:.1f}" '
                         f'stroke="#efefef" stroke-width="1"/>')
            parts.append(f'<text x="{PL-3}" y="{yp+3:.1f}" text-anchor="end" '
                         f'font-size="7" fill="#bbb" font-family="sans-serif">{yt}</text>')

    # 50% reference line (election threshold etc.)
    if 50 in range(int(ymin), int(ymax)+1):
        yp50 = py(50)
        parts.append(f'<line x1="{PL}" y1="{yp50:.1f}" x2="{W-PR}" y2="{yp50:.1f}" '
                     f'stroke="#aaa" stroke-width="1" stroke-dasharray="3,3"/>')

    color = "#e67e22" if flat else "#1a6ea8"
    pts_str = " ".join(f"{px(x):.1f},{py(y):.1f}" for x, y in points)
    parts.append(f'<polyline points="{pts_str}" fill="none" stroke="{color}" '
                 f'stroke-width="2.5" stroke-linejoin="round"/>')
    for x, y in points:
        parts.append(f'<circle cx="{px(x):.1f}" cy="{py(y):.1f}" r="3.5" fill="{color}"/>')
        # Year label (only show if not too crowded)
        if len(points) <= 7:
            parts.append(f'<text x="{px(x):.1f}" y="{H-PB+10}" text-anchor="middle" '
                         f'font-size="7" fill="#999" font-family="sans-serif">{x}</text>')
            parts.append(f'<text x="{px(x)+4:.1f}" y="{py(y)-5:.1f}" '
                         f'font-size="8" fill="{color}" font-family="sans-serif">{y}%</text>')
    return _svg_wrap(W, H, "".join(parts))


def _render_comparison_bar(left: int, left_label: str, right: int, right_label: str) -> str:
    W, H = 240, 88
    PL = 72
    BAR_H = 18
    max_v = max(left, right) * 1.15 or 1

    def bw(v): return max(3, (v / max_v) * (W - PL - 12))

    parts = [f'<rect width="{W}" height="{H}" fill="#fafafa" rx="4"/>']
    # Left / target bar
    w1 = bw(left)
    parts.append(f'<rect x="{PL}" y="14" width="{w1:.1f}" height="{BAR_H}" fill="#2d7a3a" rx="2"/>')
    parts.append(f'<text x="{PL-4}" y="27" text-anchor="end" font-size="9" fill="#555" font-family="sans-serif">{left_label}</text>')
    parts.append(f'<text x="{PL+w1+4:.1f}" y="27" font-size="10" font-weight="bold" fill="#2d7a3a" font-family="sans-serif">{left}</text>')
    # Right / actual bar
    w2 = bw(right)
    parts.append(f'<rect x="{PL}" y="48" width="{w2:.1f}" height="{BAR_H}" fill="#c0392b" rx="2"/>')
    parts.append(f'<text x="{PL-4}" y="61" text-anchor="end" font-size="9" fill="#555" font-family="sans-serif">{right_label}</text>')
    parts.append(f'<text x="{PL+w2+4:.1f}" y="61" font-size="10" font-weight="bold" fill="#c0392b" font-family="sans-serif">{right}</text>')
    return _svg_wrap(W, H, "".join(parts))


def _render_person_grid(total: int, hi: int) -> str:
    cols = 10
    rows = math.ceil(total / cols)
    cell = 20
    W, H = cols * cell + 10, rows * cell + 26
    parts = [f'<rect width="{W}" height="{H}" fill="#fafafa" rx="4"/>']
    for n in range(total):
        col, row = n % cols, n // cols
        cx = 5 + col * cell + cell // 2
        cy = 4 + row * cell + cell // 2
        color = "#c0392b" if n < hi else "#d8d8d8"
        r = cell * 0.3
        # Head
        parts.append(f'<circle cx="{cx}" cy="{cy-2}" r="{r:.1f}" fill="{color}"/>')
        # Body arc
        bx, by = cx - r * 0.7, cy + r * 0.4
        parts.append(f'<path d="M {bx:.1f} {cy+r*2:.1f} Q {cx} {cy+r*3:.1f} '
                     f'{cx+r*0.7:.1f} {cy+r*2:.1f}" fill="none" stroke="{color}" stroke-width="2"/>')
    parts.append(f'<text x="{W//2}" y="{H-4}" text-anchor="middle" font-size="9" '
                 f'font-weight="bold" fill="#c0392b" font-family="sans-serif">{hi} of {total}</text>')
    return _svg_wrap(W, H, "".join(parts))


def _render_ratio_bar(ratio: int) -> str:
    W, H = 240, 80
    PL = 72
    BAR_H = 18
    unit_w = (W - PL - 14) / ratio

    parts = [f'<rect width="{W}" height="{H}" fill="#fafafa" rx="4"/>']
    w1 = unit_w
    parts.append(f'<rect x="{PL}" y="14" width="{w1:.1f}" height="{BAR_H}" fill="#888" rx="2"/>')
    parts.append(f'<text x="{PL-4}" y="27" text-anchor="end" font-size="9" fill="#555" font-family="sans-serif">Black family</text>')
    parts.append(f'<text x="{PL+w1+4:.1f}" y="27" font-size="9" fill="#888" font-family="sans-serif">1×</text>')
    w2 = W - PL - 14
    parts.append(f'<rect x="{PL}" y="46" width="{w2}" height="{BAR_H}" fill="#1a6ea8" rx="2"/>')
    parts.append(f'<text x="{PL-4}" y="59" text-anchor="end" font-size="9" fill="#555" font-family="sans-serif">White family</text>')
    parts.append(f'<text x="{PL+w2+4}" y="59" font-size="11" font-weight="bold" fill="#1a6ea8" font-family="sans-serif">{ratio}×</text>')
    return _svg_wrap(W, H, "".join(parts))


def _render_checklist(good: list[str], bad: list[str]) -> str:
    items = [(True, g) for g in good] + [(False, b) for b in bad]
    W = 240
    line_h = 20
    H = max(50, len(items) * line_h + 16)
    parts = [f'<rect width="{W}" height="{H}" fill="#fafafa" rx="4"/>']
    for i, (ok, label) in enumerate(items):
        y = 16 + i * line_h
        icon, color = ("✓", "#2d7a3a") if ok else ("✗", "#c0392b")
        parts.append(f'<text x="10" y="{y}" font-size="12" font-weight="bold" '
                     f'fill="{color}" font-family="sans-serif">{icon}</text>')
        short = str(label)[:32]
        parts.append(f'<text x="26" y="{y}" font-size="10" fill="#333" '
                     f'font-family="sans-serif">{_esc(short)}</text>')
    return _svg_wrap(W, H, "".join(parts))


def render_data_chart(desc: str) -> str:
    """Return an inline SVG chart for a CUSTOM data-card SHOW description, or '' if not parseable."""
    data = _parse_chart_data(desc)
    if not data:
        return ""
    t = data["type"]
    if t == "gauges":
        return _render_gauges(data["values"])
    if t == "line":
        return _render_line_chart(data["points"], flat=data.get("flat", False))
    if t == "comparison_bar":
        return _render_comparison_bar(data["left"], data["left_label"],
                                      data["right"], data["right_label"])
    if t == "person_grid":
        return _render_person_grid(data["total"], data["hi"])
    if t == "ratio_bar":
        return _render_ratio_bar(data["ratio"])
    if t == "checklist":
        return _render_checklist(data["good"], data["bad"])
    return ""


# ─── HTML generation ──────────────────────────────────────────────────────────

SOURCE_COLORS = {
    "ARCHIVE":  ("#1a6ea8", "#e8f4fd", "📽"),   # blue
    "CUSTOM":   ("#7b3fa0", "#f5eeff", "🎨"),   # purple
    "SELF":     ("#555555", "#f5f5f5", "🎥"),   # grey
    "GE":       ("#2d7a3a", "#eaf7ec", "🌍"),   # green
    "LICENSED": ("#c0392b", "#fff0ee", "⚖"),    # red (AP/Getty)
    "GENERAL":  ("#1a6ea8", "#e8f4fd", "🔍"),   # blue fallback
}

# Map prod_source strings to our source_type for color
_PROD_SOURCE_MAP = {
    "AP": "LICENSED", "GETTY": "LICENSED", "REUTERS": "LICENSED",
    "AP/GETTY": "LICENSED", "AP / GETTY": "LICENSED",
    "AP/REUTERS": "LICENSED", "W": "ARCHIVE", "WIKIMEDIA": "ARCHIVE",
    "GE": "GE", "GOOGLE EARTH": "GE",
    "CUSTOM": "CUSTOM", "SELF": "SELF",
    "E": "ARCHIVE", "ENVATO": "ARCHIVE",
    "SABC": "LICENSED",
}


def _resolve_color(seg: Segment) -> tuple[str, str, str]:
    """Return (border_color, bg_color, icon)."""
    src = seg.source_type
    # Override with prod_source if available and more specific
    if seg.prod_source:
        mapped = _PROD_SOURCE_MAP.get(seg.prod_source.upper().strip())
        if mapped:
            src = mapped
    return SOURCE_COLORS.get(src, SOURCE_COLORS["GENERAL"])


HTML_CSS = """
  :root {
    --font-body: 'Georgia', 'Times New Roman', serif;
    --font-ui: -apple-system, 'Segoe UI', sans-serif;
    --max-w: 960px;
    --text: #1a1a1a;
    --muted: #666;
    --bg: #f5f4f1;
    --divider: #ddd;
    --gutter: 52px;
  }
  * { box-sizing: border-box; margin: 0; padding: 0; }
  body {
    background: var(--bg);
    color: var(--text);
    font-family: var(--font-body);
    font-size: 16px;
    line-height: 1.8;
    padding: 36px 16px 80px;
  }
  .page { max-width: var(--max-w); margin: 0 auto; }

  /* ── Script header ── */
  .script-header {
    border-bottom: 3px solid #1a1a1a;
    padding-bottom: 20px;
    margin-bottom: 32px;
  }
  .script-header h1 { font-family: var(--font-ui); font-size: 24px; font-weight: 700; }
  .script-header .meta { color: var(--muted); font-family: var(--font-ui); font-size: 12px; margin-top: 6px; }
  .script-header .stats { display: flex; gap: 18px; margin-top: 10px; flex-wrap: wrap; }
  .stat-chip { background: #1a1a1a; color: #fff; font-family: var(--font-ui); font-size: 11px; padding: 3px 9px; border-radius: 20px; }

  /* ── Legend ── */
  .legend-row { display: flex; gap: 12px; flex-wrap: wrap; font-family: var(--font-ui); font-size: 11px; margin-bottom: 28px; align-items: center; }
  .legend-row span { color: var(--muted); }
  .leg-chip { display: flex; align-items: center; gap: 4px; padding: 2px 8px; border-radius: 20px; font-weight: 600; font-size: 10px; }

  /* ── Section header ── */
  .section-header {
    background: #1a1a1a; color: #fff;
    font-family: var(--font-ui); font-size: 11px; font-weight: 700;
    letter-spacing: .14em; text-transform: uppercase;
    padding: 8px 14px; border-radius: 3px;
    margin: 36px 0 4px var(--gutter);
  }

  /* ── A/V row ── */
  .av-row {
    display: grid;
    grid-template-columns: var(--gutter) 1fr 270px;
    gap: 0;
    border-top: 1px solid #e8e6e0;
    min-height: 44px;
  }
  .av-row:last-child { border-bottom: 1px solid #e8e6e0; }

  /* ── Delivery gutter ── */
  .av-tag {
    display: flex;
    align-items: flex-start;
    padding: 10px 6px 0 0;
    justify-content: flex-end;
  }
  .dtag {
    font-family: var(--font-ui); font-size: 9px; font-weight: 800;
    letter-spacing: .1em; text-transform: uppercase;
    padding: 2px 5px; border-radius: 3px; white-space: nowrap;
  }
  .dtag-vo      { background: #e8f0fe; color: #1a5fa8; }
  .dtag-th      { background: #fde8e8; color: #b71c1c; }
  .dtag-sponsor { background: #fff8e1; color: #8d6e00; }
  .dtag-bump    { background: #e8f5e9; color: #1b5e20; }

  /* ── Script text (audio column) ── */
  .av-text {
    padding: 10px 18px 10px 12px;
    font-size: 15.5px;
    line-height: 1.75;
  }
  .av-text p { margin-bottom: 8px; }
  .av-text p:last-child { margin-bottom: 0; }

  /* ── Visual column ── */
  .av-visual {
    padding: 8px 0 8px 0;
    border-left: 1px solid #e8e6e0;
    min-width: 0;
  }
  .av-visual-empty { /* no content */ }

  /* ── Visual card (inside av-visual) ── */
  .visual-card {
    background: #fff;
    border-left: 4px solid #999;
    margin: 0 0 0 0;
    padding: 9px 10px;
    font-family: var(--font-ui);
    font-size: 12px;
  }
  .vc-top { display: flex; gap: 10px; align-items: flex-start; }
  .vc-thumb { flex-shrink: 0; }
  .vc-thumb img { display: block; border-radius: 3px; width: 100px; height: auto; max-height: 72px; object-fit: cover; }
  .vc-no-img { width: 60px; height: 44px; background: #f0f0f0; border-radius: 3px;
               display: flex; align-items: center; justify-content: center;
               font-size: 20px; color: #ccc; }
  .vc-chart { margin: 4px 0 2px; }
  .vc-body { flex: 1; min-width: 0; }
  .vc-badges { display: flex; gap: 4px; flex-wrap: wrap; margin-bottom: 4px; }
  .vc-badge { font-size: 9px; font-weight: 700; letter-spacing: .05em;
              padding: 1px 6px; border-radius: 20px; background: #eee; color: #555; }
  .vc-shot { background: #1a1a1a; color: #fff; }
  .vc-desc { font-size: 11px; color: #444; line-height: 1.45; margin-bottom: 3px; }
  .vc-attr { font-size: 10px; color: var(--muted); line-height: 1.3; }
  .vc-attr a { color: #1a6ea8; text-decoration: none; }
  .vc-attr a:hover { text-decoration: underline; }
  .vc-notes { font-size: 10px; color: #999; margin-top: 3px; font-style: italic; }

  /* ── Full-width rows (section headers, dividers, pauses) ── */
  .row-full {
    grid-column: 1 / -1;
  }

  /* ── Quote ── */
  .quote-row .av-text {
    border-left: 3px solid #ccc;
    color: #444; font-style: italic;
    background: #fafaf7;
  }

  /* ── Pause ── */
  .pause-row .av-text {
    font-family: var(--font-ui); font-size: 11px;
    letter-spacing: .12em; text-transform: uppercase;
    color: #bbb; text-align: center;
  }

  /* ── Divider ── */
  .divider-line {
    grid-column: 1 / -1;
    border: none; border-top: 2px solid #e0ddd7;
    margin: 4px 0;
  }

  /* ── Print ── */
  @media print {
    body { background: #fff; font-size: 11pt; padding: 0; }
    .visual-card { break-inside: avoid; }
    .vc-thumb img { width: 80px; }
    .av-row { break-inside: avoid; }
    .legend-row { display: none; }
  }
"""


def _esc(s: str) -> str:
    return html_lib.escape(s)


def _text_to_paras(text: str) -> str:
    """Convert a text block to HTML paragraphs. Escapes first, then applies inline markdown."""
    parts = re.split(r"\n{2,}", text.strip())
    out = []
    for p in parts:
        p = p.strip()
        if not p:
            continue
        # Escape all HTML entities first
        p = html_lib.escape(p)
        # Now apply simple inline markdown (on escaped text)
        p = re.sub(r"\*\*(.+?)\*\*", r"<strong>\1</strong>", p)
        p = re.sub(r"\*(.+?)\*", r"<em>\1</em>", p)
        out.append(f"<p>{p}</p>")
    return "\n".join(out)


def _word_count(segments: list[Segment]) -> int:
    return sum(len(s.content.split()) for s in segments if s.kind in ("text", "quote"))


def _visual_card_html(seg: Segment) -> str:
    """Render a visual card for a visual_cue segment."""
    border_color, bg_color, icon = _resolve_color(seg)

    # Badges line
    badges = (f'<span class="vc-badge" style="background:{bg_color};color:{border_color}">'
              f'{icon} {_esc(seg.source_type)}</span>')
    if seg.prod_source:
        badges += f' <span class="vc-badge" style="background:#f0f0f0;color:#555">{_esc(seg.prod_source)}</span>'
    if seg.shot_num:
        badges += f' <span class="vc-badge vc-shot">#{_esc(seg.shot_num)}</span>'

    # Image / chart area
    chart_svg = ""
    if seg.source_type == "CUSTOM":
        chart_svg = render_data_chart(seg.content)

    if seg.image.found:
        media_html = (
            f'<div class="vc-thumb">'
            f'<a href="{_esc(seg.image.commons_url)}" target="_blank">'
            f'<img src="{_esc(seg.image.thumb_url)}" alt="{_esc(seg.image.title)}" loading="lazy">'
            f'</a></div>'
        )
        attr_html = (
            f'<div class="vc-attr">'
            f'<a href="{_esc(seg.image.commons_url)}" target="_blank">'
            f'{_esc(seg.image.title[:55])}</a> · {_esc(seg.image.author[:40])} · '
            f'{_esc(seg.image.license)}</div>'
        )
    else:
        media_html = f'<div class="vc-no-img">{icon}</div>'
        if seg.source_type == "SELF":
            attr_info = "On-camera / talking head"
        elif seg.source_type == "CUSTOM":
            attr_info = "Custom motion graphic"
        elif seg.source_type == "GE":
            attr_info = "Google Earth aerial — use coords from prod. package"
        elif seg.source_type == "LICENSED":
            attr_info = "Licensed footage — AP / Getty / Reuters"
        else:
            attr_info = seg.image.error or ""
        attr_html = f'<div class="vc-attr">{_esc(attr_info)}</div>'

    notes_html = ""
    if seg.prod_notes:
        notes_html = f'<div class="vc-notes">&#9999; {_esc(seg.prod_notes)}</div>'

    inner = (
        f'<div class="vc-top">'
        f'{media_html}'
        f'<div class="vc-body">'
        f'<div class="vc-badges">{badges}</div>'
        f'<div class="vc-desc">{_esc(seg.content)}</div>'
        f'{attr_html}'
        f'{notes_html}'
        f'</div></div>'
    )
    if chart_svg:
        inner += f'<div class="vc-chart">{chart_svg}</div>'

    return (f'<div class="visual-card" '
            f'style="border-left-color:{border_color};background:{bg_color}">'
            f'{inner}</div>')


def build_html(segments: list[Segment], script_path: Path, package_path: Optional[Path]) -> str:
    # Assign delivery types before rendering
    assign_delivery_types(segments)

    title = _esc(script_path.stem.replace("_", " ").title())
    wc = _word_count(segments)
    est_mins = round(wc / 150)
    visual_count  = sum(1 for s in segments if s.kind == "visual_cue")
    found_count   = sum(1 for s in segments if s.kind == "visual_cue" and s.image.found)
    chart_count   = sum(1 for s in segments
                        if s.kind == "visual_cue" and s.source_type == "CUSTOM"
                        and render_data_chart(s.content))

    package_note = f' · Package: <em>{_esc(package_path.name)}</em>' if package_path else ""

    # ── Legend ───────────────────────────────────────────────────
    leg_items = [
        ("#1a5fa8", "#e8f0fe", "VO", "Voiceover"),
        ("#b71c1c", "#fde8e8", "TH", "Talking Head"),
        ("#8d6e00", "#fff8e1", "SP", "Sponsor"),
        ("#1b5e20", "#e8f5e9", "BU", "Subscribe Bump"),
    ]
    leg_delivery = "".join(
        f'<span class="leg-chip" style="background:{bg};color:{c}">{t}</span> {_esc(lbl)}'
        f'&nbsp;&nbsp;'
        for c, bg, t, lbl in leg_items
    )
    leg_source = "".join(
        f'<span class="leg-chip" style="background:{bg};color:{c}">{icon}</span> {label}&nbsp;&nbsp;'
        for label, (c, bg, icon) in SOURCE_COLORS.items() if label != "GENERAL"
    )

    # ── Build A/V rows ────────────────────────────────────────────
    # We pair each visual_cue with the text that FOLLOWS it (spoken over that visual).
    # Pending text is held until we know whether to attach it to a visual or emit standalone.
    rows_html: list[str] = []
    pending_text: Optional[Segment] = None   # text waiting to be paired
    pending_visual: Optional[Segment] = None # visual waiting for its text

    DELIVERY_CLASSES = {
        "VO": "dtag-vo", "TH": "dtag-th",
        "SPONSOR": "dtag-sponsor", "BUMP": "dtag-bump",
    }

    def _dtag(delivery: str) -> str:
        cls = DELIVERY_CLASSES.get(delivery, "dtag-vo")
        short = {"VO": "VO", "TH": "TH", "SPONSOR": "SP", "BUMP": "SUB"}.get(delivery, delivery)
        return f'<span class="dtag {cls}">{short}</span>'

    def _text_row(seg: Segment, visual: Optional[Segment] = None, extra_class: str = "") -> str:
        tag_html  = f'<div class="av-tag">{_dtag(seg.delivery)}</div>'
        text_html = f'<div class="av-text">{_text_to_paras(seg.content)}</div>'
        if visual:
            vis_html = f'<div class="av-visual">{_visual_card_html(visual)}</div>'
        else:
            vis_html = '<div class="av-visual av-visual-empty"></div>'
        return f'<div class="av-row {extra_class}">{tag_html}{text_html}{vis_html}</div>\n'

    def _visual_only_row(visual: Segment) -> str:
        """Visual cue with no associated text (rare, between two text-less visuals)."""
        tag_html  = '<div class="av-tag"></div>'
        text_html = '<div class="av-text"></div>'
        vis_html  = f'<div class="av-visual">{_visual_card_html(visual)}</div>'
        return f'<div class="av-row">{tag_html}{text_html}{vis_html}</div>\n'

    def flush_pending():
        nonlocal pending_text, pending_visual
        if pending_text and pending_visual:
            rows_html.append(_text_row(pending_text, pending_visual))
        elif pending_text:
            rows_html.append(_text_row(pending_text))
        elif pending_visual:
            rows_html.append(_visual_only_row(pending_visual))
        pending_text = None
        pending_visual = None

    for seg in segments:
        if seg.kind == "section_header":
            flush_pending()
            rows_html.append(f'<div class="section-header">{_esc(seg.content)}</div>\n')

        elif seg.kind == "divider":
            flush_pending()
            rows_html.append('<hr class="divider-line">\n')

        elif seg.kind == "pause":
            flush_pending()
            pause_seg = Segment(kind="text", content="— PAUSE —", delivery="VO")
            rows_html.append(_text_row(pause_seg, extra_class="pause-row"))

        elif seg.kind == "quote":
            flush_pending()
            q_seg = Segment(kind="text", content=f'"{seg.content}"', delivery="VO")
            rows_html.append(_text_row(q_seg, extra_class="quote-row"))

        elif seg.kind == "text":
            if pending_text:
                # We have two text blocks in a row (no visual between them);
                # emit the pending one standalone
                rows_html.append(_text_row(pending_text, pending_visual))
                pending_visual = None
            pending_text = seg

        elif seg.kind == "visual_cue":
            if pending_visual and not pending_text:
                # Two visuals in a row with no text between them
                rows_html.append(_visual_only_row(pending_visual))
                pending_visual = None
            if pending_text:
                # Pair the waiting text with THIS visual
                rows_html.append(_text_row(pending_text, seg))
                pending_text = None
                pending_visual = None
            else:
                pending_visual = seg

    flush_pending()
    body_html = "".join(rows_html)

    return f"""<!DOCTYPE html>
<html lang="en">
<head>
<meta charset="UTF-8">
<meta name="viewport" content="width=device-width, initial-scale=1">
<title>{title} — Production Script</title>
<style>{HTML_CSS}</style>
</head>
<body>
<div class="page">

  <div class="script-header">
    <h1>{title}</h1>
    <div class="meta">Script: <em>{_esc(script_path.name)}</em>{package_note}</div>
    <div class="stats">
      <span class="stat-chip">~{wc:,} words</span>
      <span class="stat-chip">~{est_mins} min runtime</span>
      <span class="stat-chip">{visual_count} visuals</span>
      <span class="stat-chip">{found_count} images fetched</span>
      <span class="stat-chip">{chart_count} charts rendered</span>
    </div>
  </div>

  <div class="legend-row">
    <span>Delivery:</span> {leg_delivery}
    <span style="margin-left:12px">Visual source:</span> {leg_source}
  </div>

  {body_html}

</div>
</body>
</html>"""


# ─── DOCX generation ──────────────────────────────────────────────────────────

def _set_cell_bg(cell, hex_color: str):
    """Set a table cell's background fill color."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color.lstrip("#"))
    tcPr.append(shd)


def _set_cell_border(cell, color: str, side: str = "left", size: int = 24):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = tcPr.find(qn("w:tcBorders"))
    if tcBorders is None:
        tcBorders = OxmlElement("w:tcBorders")
        tcPr.append(tcBorders)
    border = OxmlElement(f"w:{side}")
    border.set(qn("w:val"), "single")
    border.set(qn("w:sz"), str(size))
    border.set(qn("w:space"), "0")
    border.set(qn("w:color"), color.lstrip("#"))
    tcBorders.append(border)


def build_docx(segments: list[Segment], script_path: Path) -> Document:
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    # Title
    title_para = doc.add_paragraph()
    title_run = title_para.add_run(script_path.stem.replace("_", " ").replace("-", " ").title())
    title_run.bold = True
    title_run.font.size = Pt(20)
    title_para.paragraph_format.space_after = Pt(6)

    wc = _word_count(segments)
    meta_para = doc.add_paragraph(f"~{wc:,} words  ·  ~{round(wc/150)} min runtime")
    meta_para.runs[0].font.color.rgb = RGBColor(0x88, 0x88, 0x88)
    meta_para.runs[0].font.size = Pt(10)

    guide = doc.add_paragraph(
        "Edit this document in Word, then sync: "
        "python bin/build_visual_script.py SCRIPT.md --sync-docx visual_script.docx"
    )
    guide.runs[0].font.size = Pt(9)
    guide.runs[0].font.color.rgb = RGBColor(0x99, 0x99, 0x99)
    guide.runs[0].italic = True
    doc.add_paragraph()

    def add_section_header(text: str):
        p = doc.add_paragraph()
        r = p.add_run(text.upper())
        r.bold = True
        r.font.size = Pt(11)
        r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        p.paragraph_format.space_before = Pt(18)
        p.paragraph_format.space_after = Pt(6)
        # Dark background via paragraph shading
        pPr = p._p.get_or_add_pPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), "1a1a1a")
        pPr.append(shd)
        # Indent
        p.paragraph_format.left_indent = Pt(8)

    _DELIVERY_LABEL = {"VO": "[VO]", "TH": "[TH]", "SPONSOR": "[SP]", "BUMP": "[SUB]"}

    def _add_runs_with_markdown(paragraph, text: str):
        for chunk in re.split(r"(\*\*.*?\*\*|\*.*?\*)", text):
            if chunk.startswith("**") and chunk.endswith("**"):
                r = paragraph.add_run(chunk[2:-2])
                r.bold = True
            elif chunk.startswith("*") and chunk.endswith("*") and len(chunk) > 2:
                r = paragraph.add_run(chunk[1:-1])
                r.italic = True
            else:
                paragraph.add_run(chunk)

    def add_spoken_block(seg: Segment):
        """Editable narration with delivery prefix (round-trip safe)."""
        prefix = _DELIVERY_LABEL.get(seg.delivery, "[VO]")
        parts = re.split(r"\n{2,}", seg.content.strip())
        for i, part in enumerate(parts):
            part = part.strip()
            if not part:
                continue
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(6)
            if i == 0:
                label_run = p.add_run(f"{prefix} ")
                label_run.bold = True
                label_run.font.size = Pt(9)
                label_run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
            _add_runs_with_markdown(p, part)

    def add_visual_card(seg: Segment):
        border_color, bg_color, icon = _resolve_color(seg)

        # Editable SHOW line (primary round-trip marker)
        show_p = doc.add_paragraph()
        show_p.paragraph_format.space_before = Pt(8)
        show_p.paragraph_format.space_after = Pt(4)
        show_p.paragraph_format.left_indent = Inches(0.15)
        show_r = show_p.add_run(f"[SHOW: {seg.content}]")
        show_r.bold = True
        show_r.font.size = Pt(10)
        show_r.font.color.rgb = RGBColor(*[int(border_color.lstrip("#")[i:i+2], 16) for i in (0, 2, 4)])

        # Optional image table (not parsed back — regenerated on sync)
        table = doc.add_table(rows=1, cols=2)
        table.style = "Table Grid"
        table.autofit = False

        # Col 0: narrow color stripe
        stripe_cell = table.cell(0, 0)
        stripe_cell.width = Inches(0.15)
        _set_cell_bg(stripe_cell, border_color)
        stripe_cell.text = ""

        # Col 1: content
        content_cell = table.cell(0, 1)
        content_cell.width = Inches(6.2)
        _set_cell_bg(content_cell, bg_color)

        # Source badge line
        badge_para = content_cell.paragraphs[0]
        badge_para.paragraph_format.space_before = Pt(4)
        badge_r = badge_para.add_run(f"{icon}  {seg.source_type}")
        badge_r.bold = True
        badge_r.font.size = Pt(9)
        badge_r.font.color.rgb = RGBColor(*[int(border_color.lstrip("#")[i:i+2], 16) for i in (0, 2, 4)])
        if seg.prod_source:
            badge_para.add_run(f"  ·  {seg.prod_source}").font.size = Pt(9)
        if seg.shot_num:
            shot_r = badge_para.add_run(f"  SHOT {seg.shot_num}")
            shot_r.font.size = Pt(9)
            shot_r.bold = True

        # Image (if available)
        if seg.image.found:
            img_bytes = fetch_image_bytes(seg.image.thumb_url)
            if img_bytes:
                img_para = content_cell.add_paragraph()
                img_para.paragraph_format.space_before = Pt(4)
                try:
                    run = img_para.add_run()
                    run.add_picture(io.BytesIO(img_bytes), width=Inches(2.4))
                except Exception:
                    img_para.add_run(f"[Image: {seg.image.title[:60]}]")

            # Attribution
            attr_para = content_cell.add_paragraph(
                f"📷 {seg.image.title[:70]}  |  {seg.image.author[:50]}  |  {seg.image.license}"
            )
            attr_para.runs[0].font.size = Pt(8)
            attr_para.runs[0].font.color.rgb = RGBColor(0x88, 0x88, 0x88)
        else:
            if seg.prod_notes:
                notes_para = content_cell.add_paragraph(f"📝 {seg.prod_notes}")
                notes_para.runs[0].font.size = Pt(9)
                notes_para.runs[0].font.color.rgb = RGBColor(0x88, 0x88, 0x88)
                notes_para.runs[0].italic = True

        # Space after table
        doc.add_paragraph().paragraph_format.space_after = Pt(4)

    # Build document (one segment → one editable block; preserves delivery tags)
    for seg in segments:
        if seg.kind == "section_header":
            add_section_header(seg.content)

        elif seg.kind == "divider":
            doc.add_paragraph("---").paragraph_format.space_after = Pt(4)

        elif seg.kind == "pause":
            p = doc.add_paragraph("[PAUSE]")
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.runs[0].font.color.rgb = RGBColor(0xAA, 0xAA, 0xAA)
            p.runs[0].font.size = Pt(10)

        elif seg.kind == "quote":
            p = doc.add_paragraph(f'"{seg.content}"')
            p.runs[0].italic = True
            p.paragraph_format.left_indent = Inches(0.4)
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(6)

        elif seg.kind == "text":
            add_spoken_block(seg)

        elif seg.kind == "visual_cue":
            add_visual_card(seg)

    return doc


# ─── Entry point ──────────────────────────────────────────────────────────────

def main():
    parser = argparse.ArgumentParser(
        description="Generate a visual script document with inline graphics.",
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog=__doc__,
    )
    parser.add_argument("script", type=Path, help="Path to the script .md file")
    parser.add_argument("--package", "-p", type=Path, default=None,
                        help="Path to production package .md (optional, improves source matching)")
    parser.add_argument("--no-fetch", action="store_true",
                        help="Skip Wikimedia image fetching (faster, offline)")
    parser.add_argument("--html-only", action="store_true",
                        help="Only generate HTML, skip DOCX")
    parser.add_argument("--docx-only", action="store_true",
                        help="Only generate DOCX, skip HTML")
    parser.add_argument("--context", "-c", default="",
                        help="Topic context hint appended to image searches (e.g. 'South Africa')")
    parser.add_argument("--sync-docx", type=Path, default=None, metavar="DOCX",
                        help="Read edited visual_script.docx, update the .md script, regenerate HTML/DOCX")
    args = parser.parse_args()

    script_path: Path = args.script.resolve()
    if not script_path.exists():
        print(f"Error: script not found: {script_path}", file=sys.stderr)
        sys.exit(1)

    package_path: Optional[Path] = args.package
    if package_path:
        package_path = package_path.resolve()
        if not package_path.exists():
            print(f"Warning: package not found: {package_path}", file=sys.stderr)
            package_path = None

    out_dir = script_path.parent
    html_out = out_dir / "visual_script.html"
    docx_out = out_dir / "visual_script.docx"

    if args.sync_docx:
        docx_path = args.sync_docx.resolve()
        if not docx_path.exists():
            print(f"Error: DOCX not found: {docx_path}", file=sys.stderr)
            sys.exit(1)
        sync_from_docx(
            docx_path,
            script_path,
            package_path,
            no_fetch=args.no_fetch,
            html_only=args.html_only,
            docx_only=args.docx_only,
            context=args.context,
        )
        return

    print(f"\n=== Visual Script Builder ===")
    print(f"Script : {script_path.name}")
    if package_path:
        print(f"Package: {package_path.name}")
    print()

    # Parse
    print("Parsing script...", end=" ")
    segments = parse_script(script_path)
    visual_count = sum(1 for s in segments if s.kind == "visual_cue")
    print(f"OK  {len(segments)} segments, {visual_count} visual cues")

    # Assign delivery types immediately after parsing
    assign_delivery_types(segments)

    # Match production package
    if package_path:
        print("Matching production package...", end=" ")
        rows = parse_production_package(package_path)
        match_visuals_to_package(segments, rows)
        matched = sum(1 for s in segments if s.kind == "visual_cue" and s.prod_source)
        print(f"OK  {matched}/{visual_count} visuals matched to package rows")

    # Fetch images
    if not args.no_fetch:
        print("\nFetching images from Wikipedia/Commons...")
        # Derive a context hint from the script filename (e.g. "south_africa" → "South Africa")
        stem = script_path.stem.replace("_", " ").replace("-", " ")
        context_hint = args.context or stem[:30]
        fetch_images_for_segments(segments, verbose=True, context_hint=context_hint)
        found = sum(1 for s in segments if s.kind == "visual_cue" and s.image.found)
        print(f"\n  {found}/{visual_count} images found")

    # HTML output
    if not args.docx_only:
        print("\nBuilding HTML...", end=" ")
        html_content = build_html(segments, script_path, package_path)
        html_out.write_text(html_content, encoding="utf-8")
        print(f"OK  {html_out.name} ({len(html_content) // 1024} KB)")

    # DOCX output
    if not args.html_only:
        print("Building DOCX...", end=" ")
        doc = build_docx(segments, script_path)
        doc.save(str(docx_out))
        print(f"OK  {docx_out.name}")

    print(f"\nDone! Outputs in: {out_dir}")
    if not args.docx_only:
        print(f"  HTML : {html_out}")
    if not args.html_only:
        print(f"  DOCX : {docx_out}")


if __name__ == "__main__":
    main()
