#!/usr/bin/env python3
"""Build a TH-only teleprompter DOCX from a visual script markdown file.

Extracts spoken blocks that follow a 🎥 SELF marker (talking head cues).
Output is formatted for phone teleprompter apps: large type, filming order, clean text.

Usage:
  python bin/build_th_teleprompter.py "generatedScripts/SA Captured Rainbow/dark_side_of_south_africa_script.md"
"""

from __future__ import annotations

import argparse
import re
import sys
from pathlib import Path

# Reuse script body parsing from visual script builder
sys.path.insert(0, str(Path(__file__).resolve().parent))
from build_visual_script import SECTION_RE, _find_script_body_range  # noqa: E402

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_LINE_SPACING

SELF_MARKER = re.compile(r"^\s*🎥\s*SELF\s*$")
SHOW_RE = re.compile(r"\[SHOW:[^\]]+\]", re.IGNORECASE)


def extract_th_blocks(path: Path) -> list[dict]:
    lines = path.read_text(encoding="utf-8").splitlines()
    start, end = _find_script_body_range(lines)
    body = lines[start:end]

    blocks: list[dict] = []
    pending_self = False
    buffer: list[str] = []
    current_section = ""

    def flush():
        nonlocal buffer, pending_self
        if not pending_self or not buffer:
            buffer = []
            pending_self = False
            return
        text = "\n\n".join(buffer).strip()
        if text:
            blocks.append({"section": current_section, "text": text})
        buffer = []
        pending_self = False

    for line in body:
        stripped = line.strip()

        if SECTION_RE.match(stripped):
            flush()
            m = SECTION_RE.match(stripped)
            current_section = m.group(1) or m.group(2) or ""
            continue

        if SELF_MARKER.match(line):
            flush()
            pending_self = True
            continue

        if SHOW_RE.search(stripped) or stripped.startswith("[PAUSE]"):
            flush()
            continue

        if not stripped:
            continue

        if pending_self:
            buffer.append(stripped)

    flush()
    return blocks


def build_teleprompter_doc(blocks: list[dict], script_path: Path, out_path: Path) -> None:
    doc = Document()
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(22)
    normal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    normal.paragraph_format.line_spacing = 1.35
    normal.paragraph_format.space_after = Pt(12)

    title = doc.add_paragraph()
    tr = title.add_run("Talking Head — Teleprompter")
    tr.bold = True
    tr.font.size = Pt(14)
    tr.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    sub = doc.add_paragraph()
    sr = sub.add_run(script_path.stem.replace("_", " "))
    sr.font.size = Pt(12)
    sr.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

    doc.add_paragraph()

    note = doc.add_paragraph()
    nr = note.add_run(f"{len(blocks)} clips · film in order · same setup for all")
    nr.font.size = Pt(11)
    nr.italic = True
    nr.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

    doc.add_paragraph()

    for i, block in enumerate(blocks, start=1):
        if i > 1:
            sep = doc.add_paragraph()
            sep.paragraph_format.space_before = Pt(18)
            sep.paragraph_format.space_after = Pt(6)
            sr = sep.add_run("— — —")
            sr.font.color.rgb = RGBColor(0xCC, 0xCC, 0xCC)
            sr.font.size = Pt(11)

        hdr = doc.add_paragraph()
        hr = hdr.add_run(f"CLIP {i}")
        hr.bold = True
        hr.font.size = Pt(12)
        hr.font.color.rgb = RGBColor(0x44, 0x44, 0x44)
        if block["section"]:
            hr2 = hdr.add_run(f"  ·  {block['section']}")
            hr2.font.size = Pt(10)
            hr2.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

        for para in block["text"].split("\n\n"):
            para = para.strip()
            if not para:
                continue
            p = doc.add_paragraph(para)
            p.paragraph_format.space_after = Pt(14)

    doc.save(str(out_path))


def main() -> None:
    parser = argparse.ArgumentParser(description="Build TH-only teleprompter DOCX")
    parser.add_argument("script", type=Path, help="Script .md path")
    parser.add_argument(
        "-o", "--output", type=Path, default=None,
        help="Output .docx (default: <script_stem>_TH_teleprompter.docx beside script)",
    )
    args = parser.parse_args()

    script_path = args.script.resolve()
    if not script_path.exists():
        print(f"Error: not found: {script_path}", file=sys.stderr)
        sys.exit(1)

    out_path = args.output or script_path.with_name(f"{script_path.stem}_TH_teleprompter.docx")
    out_path = out_path.resolve()

    blocks = extract_th_blocks(script_path)
    if not blocks:
        print("Error: no TH blocks found (look for 🎥 SELF markers)", file=sys.stderr)
        sys.exit(1)

    build_teleprompter_doc(blocks, script_path, out_path)
    print(f"OK  {len(blocks)} TH clips -> {out_path.name}")
    print(f"    {out_path}")


if __name__ == "__main__":
    main()
