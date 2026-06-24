#!/usr/bin/env python3
"""
Generate data charts for the Male Loneliness Gaming script and embed them
into the visual_script.docx produced by build_visual_script.py.

Usage:
    python bin/build_loneliness_charts.py

Outputs:
    generatedScripts/Male Loneliness Epidemic/charts/  (PNG files)
    generatedScripts/Male Loneliness Epidemic/visual_script.docx  (updated in place)
"""

import io
import re
import shutil
from pathlib import Path

import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt
import matplotlib.patches as mpatches
import numpy as np
from docx import Document
from docx.shared import Inches, Pt

# ── Paths ─────────────────────────────────────────────────────────────────────
SCRIPT_DIR = Path(__file__).parent.parent / "generatedScripts" / "Male Loneliness Epidemic"
CHARTS_DIR = SCRIPT_DIR / "charts"
DOCX_PATH  = SCRIPT_DIR / "visual_script.docx"
CHARTS_DIR.mkdir(exist_ok=True)

# ── Shared style ───────────────────────────────────────────────────────────────
BRAND_DARK   = "#1a1a1a"
BRAND_BLUE   = "#1a6ea8"
BRAND_PURPLE = "#7b3fa0"
BRAND_ORANGE = "#e67e22"
BRAND_RED    = "#c0392b"
BRAND_GREEN  = "#2d7a3a"
FONT         = "DejaVu Sans"


def _save(fig, name: str) -> Path:
    p = CHARTS_DIR / name
    fig.savefig(p, dpi=150, bbox_inches="tight", facecolor=fig.get_facecolor())
    plt.close(fig)
    print(f"  Generated: {name}")
    return p


# ── Chart 1: Stat card — Male suicide rates ~20% ─────────────────────────────
def chart_suicide_stat() -> Path:
    fig, ax = plt.subplots(figsize=(6, 2.4), facecolor=BRAND_DARK)
    ax.set_facecolor(BRAND_DARK)
    ax.axis("off")
    ax.text(0.5, 0.72, "~20%", ha="center", va="center",
            fontsize=64, fontweight="bold", color="#e74c3c", fontfamily=FONT,
            transform=ax.transAxes)
    ax.text(0.5, 0.30, "increase in male suicide rates", ha="center", va="center",
            fontsize=14, color="#cccccc", fontfamily=FONT, transform=ax.transAxes)
    ax.text(0.5, 0.08, "Source: Men's Health Network / PCORI Expert Panel Report, Oct 2019",
            ha="center", va="center", fontsize=7, color="#888888",
            fontfamily=FONT, transform=ax.transAxes)
    return _save(fig, "chart_suicide_stat.png")


# ── Chart 2: Bar — Zero close friends 1990 vs 2021 ───────────────────────────
def chart_zero_friends() -> Path:
    fig, ax = plt.subplots(figsize=(6, 3.6), facecolor="white")
    ax.set_facecolor("#f9f9f9")

    years  = ["1990", "2021"]
    values = [3, 15]
    colors = [BRAND_BLUE, BRAND_RED]

    bars = ax.bar(years, values, color=colors, width=0.45, zorder=3)
    for bar, val in zip(bars, values):
        ax.text(bar.get_x() + bar.get_width() / 2, bar.get_height() + 0.4,
                f"{val}%", ha="center", va="bottom",
                fontsize=18, fontweight="bold", color=BRAND_DARK, fontfamily=FONT)

    ax.set_ylim(0, 20)
    ax.set_ylabel("Men with zero close friends (%)", fontsize=10,
                  color="#444", fontfamily=FONT)
    ax.set_title("The Friendship Recession", fontsize=13, fontweight="bold",
                 color=BRAND_DARK, fontfamily=FONT, pad=10)
    ax.yaxis.grid(True, color="#dddddd", zorder=0)
    ax.set_axisbelow(True)
    ax.spines[["top", "right", "left"]].set_visible(False)
    ax.tick_params(axis="x", labelsize=13)
    ax.tick_params(axis="y", labelsize=9, colors="#888")

    ax.text(0.5, -0.18,
            "Source: Survey Center on American Life, AEI (May 2021)",
            ha="center", transform=ax.transAxes, fontsize=7, color="#888", fontfamily=FONT)
    fig.tight_layout()
    return _save(fig, "chart_zero_friends.png")


# ── Chart 3: Grouped bar — Depression & anxiety by gaming group ───────────────
def chart_depression_anxiety() -> Path:
    groups   = ["Non-gamers", "Low-risk gamers", "High-risk gamers"]
    depression = [6.2, 4.5, 9.5]
    anxiety    = [8.6, 9.1, 14.5]

    x   = np.arange(len(groups))
    w   = 0.35
    fig, ax = plt.subplots(figsize=(7, 4), facecolor="white")
    ax.set_facecolor("#f9f9f9")

    b1 = ax.bar(x - w/2, depression, w, label="Depressive disorder",
                color=BRAND_BLUE, zorder=3)
    b2 = ax.bar(x + w/2, anxiety,    w, label="Anxiety disorder",
                color=BRAND_ORANGE, zorder=3)

    for bar, val in zip(list(b1) + list(b2), depression + anxiety):
        ax.text(bar.get_x() + bar.get_width() / 2, bar.get_height() + 0.2,
                f"{val}%", ha="center", va="bottom",
                fontsize=9, fontweight="bold", color=BRAND_DARK, fontfamily=FONT)

    ax.set_xticks(x)
    ax.set_xticklabels(groups, fontsize=10, fontfamily=FONT)
    ax.set_ylabel("Prevalence (%)", fontsize=10, color="#444", fontfamily=FONT)
    ax.set_title("Psychiatric disorder rates by gaming group", fontsize=12,
                 fontweight="bold", color=BRAND_DARK, fontfamily=FONT, pad=10)
    ax.set_ylim(0, 18)
    ax.yaxis.grid(True, color="#dddddd", zorder=0)
    ax.set_axisbelow(True)
    ax.spines[["top", "right", "left"]].set_visible(False)
    ax.legend(fontsize=9, framealpha=0.5)
    ax.tick_params(axis="y", labelsize=9, colors="#888")

    ax.text(0.5, -0.18,
            "Source: Jung et al., Psychiatry Investigation 2025 (N=5,511 Korean adults)",
            ha="center", transform=ax.transAxes, fontsize=7, color="#888", fontfamily=FONT)
    fig.tight_layout()
    return _save(fig, "chart_depression_anxiety.png")


# ── Chart 4: Grouped bar — Loneliness scores by gaming group (Korean study) ───
def chart_loneliness_scores() -> Path:
    groups     = ["Non-gamers", "Low-risk gamers", "High-risk gamers"]
    all_resp   = [5.71, 4.97, 5.40]
    male_resp  = [5.86, 4.86, 5.59]

    x  = np.arange(len(groups))
    w  = 0.35
    fig, ax = plt.subplots(figsize=(7, 4), facecolor="white")
    ax.set_facecolor("#f9f9f9")

    b1 = ax.bar(x - w/2, all_resp,  w, label="All respondents",
                color=BRAND_BLUE, zorder=3)
    b2 = ax.bar(x + w/2, male_resp, w, label="Male respondents",
                color=BRAND_GREEN, zorder=3)

    for bar, val in zip(list(b1) + list(b2), all_resp + male_resp):
        ax.text(bar.get_x() + bar.get_width() / 2, bar.get_height() + 0.04,
                f"{val:.2f}", ha="center", va="bottom",
                fontsize=9, fontweight="bold", color=BRAND_DARK, fontfamily=FONT)

    ax.set_xticks(x)
    ax.set_xticklabels(groups, fontsize=10, fontfamily=FONT)
    ax.set_ylabel("Loneliness score (LSIS-6, lower = less lonely)",
                  fontsize=9, color="#444", fontfamily=FONT)
    ax.set_title("Loneliness scores by gaming group\n(low-risk gamers score lowest)",
                 fontsize=12, fontweight="bold", color=BRAND_DARK, fontfamily=FONT, pad=10)
    ax.set_ylim(4, 6.6)
    ax.yaxis.grid(True, color="#dddddd", zorder=0)
    ax.set_axisbelow(True)
    ax.spines[["top", "right", "left"]].set_visible(False)
    ax.legend(fontsize=9, framealpha=0.5)
    ax.tick_params(axis="y", labelsize=9, colors="#888")

    ax.text(0.5, -0.18,
            "Source: Jung et al., Psychiatry Investigation 2025 — doi.org/10.30773/pi.2023.0385",
            ha="center", transform=ax.transAxes, fontsize=7, color="#888", fontfamily=FONT)
    fig.tight_layout()
    return _save(fig, "chart_loneliness_scores.png")


# ── Chart 5: Quote card — Surgeon General ─────────────────────────────────────
def chart_surgeon_general() -> Path:
    fig, ax = plt.subplots(figsize=(7, 2.8), facecolor=BRAND_DARK)
    ax.set_facecolor(BRAND_DARK)
    ax.axis("off")

    quote = ('"Our epidemic of loneliness and isolation\n'
             'has been more than half a century in the making."')
    ax.text(0.5, 0.62, quote, ha="center", va="center",
            fontsize=13, color="white", fontfamily=FONT,
            fontstyle="italic", transform=ax.transAxes,
            linespacing=1.6, multialignment="center")
    ax.text(0.5, 0.18, "— U.S. Surgeon General, Advisory on Social Connection (2023)",
            ha="center", va="center", fontsize=9, color="#aaaaaa",
            fontfamily=FONT, transform=ax.transAxes)

    # Thin top accent line
    fig.add_artist(plt.Line2D([0.08, 0.92], [0.95, 0.95],
                              transform=fig.transFigure,
                              color=BRAND_BLUE, linewidth=2))
    return _save(fig, "chart_surgeon_general.png")


# ── DOCX embedding ─────────────────────────────────────────────────────────────
# Map: substring to match in [SHOW: ...] text  →  chart PNG path
CHART_MAP = {
    "Male suicide rates up": "chart_suicide_stat.png",
    "Men with zero close friends": "chart_zero_friends.png",
    "High-risk gamers: depression 9.5%": "chart_depression_anxiety.png",
    "ONE GRAPH: Korean": "chart_loneliness_scores.png",
    "Surgeon General advisory": "chart_surgeon_general.png",
}


def embed_charts(docx_path: Path, charts_dir: Path):
    """Open the DOCX and inject chart PNGs into matching CUSTOM visual cards."""
    doc = Document(str(docx_path))

    # Build an index: paragraph index → chart file, by scanning for SHOW: text
    para_chart = {}
    for i, para in enumerate(doc.paragraphs):
        text = para.text
        if "[SHOW:" not in text:
            continue
        for key, fname in CHART_MAP.items():
            if key.lower() in text.lower():
                chart_file = charts_dir / fname
                if chart_file.exists():
                    para_chart[i] = chart_file
                break

    if not para_chart:
        print("  No matching SHOW paragraphs found — check DOCX structure.")
        return

    # Walk through document body elements, find matching paragraphs and insert image
    # into the table cell that immediately follows (the visual card table)
    body = doc.element.body
    body_children = list(body)

    for i, para in enumerate(doc.paragraphs):
        if i not in para_chart:
            continue
        chart_file = para_chart[i]

        # Find the paragraph element in body_children
        para_elem = para._element
        try:
            idx = body_children.index(para_elem)
        except ValueError:
            continue

        # The visual card table is the next sibling element
        if idx + 1 >= len(body_children):
            continue
        next_elem = body_children[idx + 1]

        # Find the content cell (col 1) of the table
        tbl_cells = next_elem.findall(".//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}tc")
        if len(tbl_cells) < 2:
            continue
        content_cell = tbl_cells[1]

        # Add a new paragraph with the image inside the content cell
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn

        # Create a new paragraph element
        new_p = OxmlElement("w:p")
        new_r = OxmlElement("w:r")
        new_drawing = OxmlElement("w:drawing")

        # Use python-docx to add a paragraph to a cell via the high-level API
        # We'll work with the cell object
        from docx.table import _Cell
        cell_obj = _Cell(content_cell, doc.tables[0])  # placeholder, overridden below

        # Direct approach: add paragraph to the cell element
        p_elem = OxmlElement("w:p")
        r_elem = OxmlElement("w:r")

        # Build inline image using python-docx's internal picture mechanism
        img_buf = io.BytesIO(chart_file.read_bytes())

        # Add a temporary paragraph to get the correct document part
        temp_para = doc.add_paragraph()
        run = temp_para.add_run()
        run.add_picture(str(chart_file), width=Inches(5.5))

        # Move the picture run's drawing element into our cell
        drawing_elem = run._r.find(
            "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}drawing"
        )
        if drawing_elem is None:
            # Try alternate namespace
            for child in run._r:
                if "drawing" in child.tag:
                    drawing_elem = child
                    break

        if drawing_elem is not None:
            r_elem.append(drawing_elem)
            p_elem.append(r_elem)
            content_cell.append(p_elem)

        # Remove the temporary paragraph
        body.remove(temp_para._element)

        safe_text = para.text[:60].encode("ascii", "replace").decode("ascii")
        print(f"  Embedded: {chart_file.name} -> matched '{safe_text}...'")

    # Back up original and save
    backup = docx_path.with_suffix(".backup.docx")
    shutil.copy2(docx_path, backup)
    doc.save(str(docx_path))
    print(f"  Saved:  {docx_path.name}  (backup: {backup.name})")


# ── Main ───────────────────────────────────────────────────────────────────────
if __name__ == "__main__":
    print("\n=== Generating charts ===")
    chart_suicide_stat()
    chart_zero_friends()
    chart_depression_anxiety()
    chart_loneliness_scores()
    chart_surgeon_general()

    print(f"\n=== Embedding charts into {DOCX_PATH.name} ===")
    if not DOCX_PATH.exists():
        print(f"  ERROR: {DOCX_PATH} not found. Run build_visual_script.py first.")
    else:
        embed_charts(DOCX_PATH, CHARTS_DIR)

    print("\nDone.")
